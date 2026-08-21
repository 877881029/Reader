from html import escape
from pathlib import Path

from docx import Document

from reader.preview.result import PreviewResult


def to_html(path: Path) -> PreviewResult:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = escape(p.text)
        if not text:
            continue
        style = (p.style.name or "") if p.style is not None else ""
        if style.startswith("Heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            parts.append(f"<h{level}>{text}</h{level}>")
        else:
            parts.append(f"<p>{text}</p>")
    for table in doc.tables:
        rows = ["<table>"]
        for row in table.rows:
            cells = "".join(f"<td>{escape(cell.text)}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        rows.append("</table>")
        parts.append("".join(rows))
    body = "\n".join(parts)
    html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{body}</body></html>"
    return PreviewResult(html=html, status_label="内置预览", kind="html")
