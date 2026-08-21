from pathlib import Path

from docx import Document

from reader.formats.docx import to_html


def test_docx_includes_heading_paragraph_and_table(tmp_path: Path):
    path = tmp_path / "s.docx"
    doc = Document()
    doc.add_heading("Spec Title", level=1)
    doc.add_paragraph("Body sentence.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    doc.save(path)
    result = to_html(path)
    assert result.status_label == "内置预览"
    assert "Spec Title" in result.html
    assert "Body sentence." in result.html
    assert "H1" in result.html
    assert "<table" in result.html.lower()


def test_docx_escapes_dangerous_characters_in_paragraph_and_table(tmp_path: Path):
    path = tmp_path / "escape.docx"
    doc = Document()
    doc.add_paragraph('<script>alert(1)</script> & "quotes"')
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "<img onerror=alert(1)>"
    doc.save(path)
    result = to_html(path)
    html = result.html
    assert "<script>" not in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
    assert "&amp;" in html
    assert "&lt;img onerror=alert(1)&gt;" in html
